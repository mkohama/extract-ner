"""
python-docx を直接使用した Word ローダー

特徴:
- python-docx のみ使用
- Markdown形式での出力(見出し、強調、リスト、リンク対応)
- テキスト、段落、表、画像を抽出
- LangChain Document対応
"""

from pathlib import Path
from typing import Any, List
from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain_core.documents import Document
from src.logger import log


class WordToMarkdownLoader:
    """
    拡張Wordファイルローダー（Unstructured代替・高機能実装）

    Unstructuredから着想を得た以下の機能を実装:
    - スタイルベースの見出し検出（Heading 1-9 → Markdown見出し）
    - リスト項目の検出（List系スタイル → Markdown箇条書き）
    - ハイパーリンクの保持（Markdownリンク形式）
    - 強調テキストの保持（太字・イタリック → Markdown形式）
    - 画像の位置記録
    """

    # 見出しスタイルのマッピング
    HEADING_STYLES = {
        "Heading 1": 1,
        "Heading 2": 2,
        "Heading 3": 3,
        "Heading 4": 4,
        "Heading 5": 5,
        "Heading 6": 6,
        "Heading 7": 7,
        "Heading 8": 8,
        "Heading 9": 9,
        "Title": 1,
        "Subtitle": 2,
    }

    # リストスタイル
    LIST_STYLES = [
        "List",
        "List 2",
        "List 3",
        "List Bullet",
        "List Bullet 2",
        "List Bullet 3",
        "List Continue",
        "List Continue 2",
        "List Continue 3",
        "List Number",
        "List Number 2",
        "List Number 3",
        "List Paragraph",
    ]

    def __init__(
        self,
        file_path: str | Path,
        extract_tables: bool = True,
        extract_images: bool = True,
        preserve_formatting: bool = True,
    ):
        """
        初期化

        Args:
            file_path: Wordファイルのパス (.docx)
            extract_tables: 表を抽出するかどうか
            extract_images: 画像の位置を記録するかどうか
            preserve_formatting: 強調テキスト（太字・イタリック）を保持するかどうか
        """
        self.file_path = Path(file_path)
        self.extract_tables = extract_tables
        self.extract_images = extract_images
        self.preserve_formatting = preserve_formatting

    def __call__(self, file_path: str) -> "WordToMarkdownLoader":
        """LangChain互換のため"""
        return WordToMarkdownLoader(
            file_path,
            self.extract_tables,
            self.extract_images,
            self.preserve_formatting,
        )

    def load(self) -> List[Document]:
        """
        Wordファイルを読み込み、Document のリストとして返す

        Returns:
            Document のリスト（通常は1つ）
        """
        try:
            doc = DocxDocument(str(self.file_path))
            content_parts = []

            # 段落と表を順番に処理
            for element in doc.element.body:
                # 段落の処理
                if element.tag.endswith("p"):
                    para = Paragraph(element, doc)
                    para_text = self._process_paragraph(para)
                    if para_text:
                        content_parts.append(para_text)

                # 表の処理
                elif element.tag.endswith("tbl") and self.extract_tables:
                    table = Table(element, doc)
                    table_text = self._extract_table(table)
                    if table_text:
                        content_parts.append(table_text)

            # 全コンテンツを結合
            full_content = "\n\n".join(content_parts)

            if not full_content.strip():
                log(f"Word file {self.file_path} has no readable content")
                return []

            # LangChain Document として返す
            document = Document(
                page_content=full_content,
                metadata={
                    "source": str(self.file_path),
                    "file_type": "docx",
                },
            )

            return [document]

        except FileNotFoundError:
            log(f"Error: File not found at {self.file_path}")
            raise  # 例外を再raise

        except Exception as e:
            log(f"Error processing Word file {self.file_path}: {e}")
            raise RuntimeError(f"Failed to process Word file: {e}") from e

    def _process_paragraph(self, paragraph: Paragraph) -> str:
        """
        段落を処理してMarkdown形式で返す

        - 見出しスタイル → Markdown見出し (# ## ###)
        - リストスタイル → Markdown箇条書き (- または 1.)
        - ハイパーリンク → Markdownリンク [text](url)
        - 強調テキスト → Markdown強調 (**bold**, *italic*)
        - 画像 → [画像]
        """
        # 空の段落をスキップ
        if not paragraph.text.strip():
            return ""

        # スタイル名を取得
        style_name = (paragraph.style and paragraph.style.name) or "Normal"

        # 見出しの処理
        if style_name in self.HEADING_STYLES:
            level = self.HEADING_STYLES[style_name]
            heading_prefix = "#" * level
            text = self._extract_paragraph_text(paragraph)
            return f"{heading_prefix} {text}"

        # リスト項目の処理
        if style_name in self.LIST_STYLES:
            text = self._extract_paragraph_text(paragraph)
            # 番号付きリストか箇条書きかの判定
            is_numbered = "Number" in style_name
            prefix = "1." if is_numbered else "-"

            # インデントレベルの取得（List 2, List 3等）
            indent_level = self._get_list_indent_level(style_name)
            indent = "  " * indent_level

            return f"{indent}{prefix} {text}"

        # 通常の段落
        text = self._extract_paragraph_text(paragraph)
        return text

    def _extract_paragraph_text(self, paragraph: Paragraph) -> str:
        """
        段落からテキストを抽出し、リンクと強調を保持
        """
        result_parts = []

        # paragraph.runs を使用して、各ランの書式を確認
        for run in paragraph.runs:
            text = run.text

            if not text:
                continue

            # 画像の検出
            if self.extract_images and self._run_contains_image(run):
                result_parts.append("[画像]")
                continue

            # 強調テキストの処理
            if self.preserve_formatting:
                # 太字とイタリックの判定
                is_bold = run.bold
                is_italic = run.italic

                if is_bold and is_italic:
                    text = f"***{text}***"
                elif is_bold:
                    text = f"**{text}**"
                elif is_italic:
                    text = f"*{text}*"

            result_parts.append(text)

        # ハイパーリンクの処理
        # python-docxのHyperlinkサポートは限定的なため、
        # paragraph.textを使用してURLパターンを検出する簡易実装
        full_text = "".join(result_parts)

        # ハイパーリンクがある場合の処理（python-docx 0.8.11+）
        if hasattr(paragraph, "hyperlinks") and paragraph.hyperlinks:
            # 既存のハイパーリンクをMarkdown形式に変換
            for hyperlink in paragraph.hyperlinks:
                if hyperlink.url:
                    link_text = hyperlink.text or hyperlink.url
                    markdown_link = f"[{link_text}]({hyperlink.url})"
                    full_text = full_text.replace(hyperlink.text, markdown_link)

        return full_text

    def _run_contains_image(self, run: Any) -> bool:
        """ランに画像が含まれているか判定"""
        try:
            # ランのXML要素に画像要素（drawing）があるか確認
            return len(run._element.xpath(".//w:drawing")) > 0
        except Exception:
            return False

    def _get_list_indent_level(self, style_name: str) -> int:
        """リストスタイル名からインデントレベルを取得"""
        # "List 2" → 1, "List Bullet 3" → 2 等
        if "2" in style_name:
            return 1
        elif "3" in style_name:
            return 2
        return 0

    def _extract_table(self, table: Table) -> str:
        """
        表をMarkdown形式で抽出

        Args:
            table: python-docx の Table オブジェクト

        Returns:
            Markdown形式の表テキスト
        """
        if not table.rows:
            return ""

        lines = []

        # ヘッダー行
        header_cells = [self._extract_cell_text(cell) for cell in table.rows[0].cells]
        if header_cells:
            lines.append("| " + " | ".join(header_cells) + " |")
            lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

        # データ行
        for row in table.rows[1:]:
            cells = [self._extract_cell_text(cell) for cell in row.cells]
            if any(cells):  # 空行をスキップ
                lines.append("| " + " | ".join(cells) + " |")

        return "\n".join(lines) if lines else ""

    def _extract_cell_text(self, cell: Any) -> str:
        """セルからテキストを抽出（複数段落対応）"""
        # セル内に複数の段落がある場合があるため、全て結合
        texts = []
        for paragraph in cell.paragraphs:
            text = paragraph.text.strip()
            if text:
                texts.append(text)
        return " ".join(texts)
